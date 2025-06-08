from multiprocessing import Value
import sqlite3
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
import requests
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import font as tkfont
import os

# Переменные для темы
DARK_THEME = {
    'bg': '#2d2d2d',
    'fg': '#ffffff',
    'entry_bg': '#3d3d3d',
    'entry_fg': '#ffffff',
    'button_bg': '#4d4d4d',
    'button_fg': '#ffffff',
    'tree_bg': '#3d3d3d',
    'tree_fg': '#ffffff',
    'tree_heading_bg': '#2d2d2d',
    'tree_heading_fg': '#ffffff',
    'frame_bg': '#2d2d2d',
    'label_frame_bg': '#2d2d2d',  
    'label_frame_fg': '#ffffff',
    'label_bg': '#2d2d2d',
    'label_fg': '#ffffff',
    'combobox_bg': '#3d3d3d',
    'combobox_fg': '#000000',
    'combobox_field_bg': '#000000',
    'combobox_arrow_color': '#ffffff'

}

LIGHT_THEME = {
    'bg': '#f0f0f0',
    'fg': '#000000',
    'entry_bg': '#ffffff',
    'entry_fg': '#000000',
    'button_bg': '#e0e0e0',
    'button_fg': '#000000',
    'tree_bg': '#ffffff',
    'tree_fg': '#000000',
    'tree_heading_bg': '#e0e0e0',
    'tree_heading_fg': '#000000',
    'frame_bg': '#f0f0f0',
    'label_frame_bg': '#ffffff', 
    'label_frame_fg': '#000000',
    'label_bg': '#f0f0f0',
    'label_fg': '#000000',
    'combobox_bg': '#ffffff',
    'combobox_fg': '#000000',
    'combobox_field_bg': '#ffffff',
    'combobox_arrow_color': '#000000'
}

current_theme = LIGHT_THEME

# === Функции для работы с настройками ===
def apply_theme():
    """Применяет выбранную тему ко всем элементам интерфейса"""
    theme = current_theme
    
    # Общие настройки
    root.config(bg=theme['bg'])
    
    # Применяем тему ко всем виджетам
    for widget in root.winfo_children():
        apply_theme_to_widget(widget, theme)
    
    # Специальные настройки для Treeview
    style = ttk.Style()
    style.theme_use('clam')
    
    style.configure("Treeview",
        background=theme['tree_bg'],
        foreground=theme['tree_fg'],
        fieldbackground=theme['tree_bg'])
    
    style.configure("Treeview.Heading",
        background=theme['tree_heading_bg'],
        foreground=theme['tree_heading_fg'])
    
    style.map('Treeview', 
        background=[('selected', '#347083')],
        foreground=[('selected', 'white')])
    
    style.configure('TCombobox',
        background=theme['combobox_bg'],
        foreground=theme['combobox_fg'],
        fieldbackground=theme['combobox_field_bg'],
        arrowcolor=theme['combobox_arrow_color'])

    style.configure('TFrame', background=theme['frame_bg'])
    style.configure('TNotebook', background=theme['bg'])
    style.configure('TNotebook.Tab', background=theme['frame_bg'], foreground=theme['fg'])

    style.map('TNotebook.Tab',
          background=[('selected', theme['bg'])],
          foreground=[('selected', theme['fg'])])
    
    update_currency_label()

def apply_theme_to_widget(widget, theme):
    """Рекурсивно применяет тему к виджету и его дочерним элементам"""
    widget_type = widget.winfo_class()
    
    try:
        if widget_type == 'TFrame' or widget_type == 'Frame':
            widget.config(bg=theme['frame_bg'])
        elif widget_type == 'Label':
            widget.config(bg=theme['label_bg'], fg=theme['label_fg'])
        elif widget_type == 'Button':
            widget.config(bg=theme['button_bg'], fg=theme['button_fg'],
                         activebackground=theme['button_bg'], activeforeground=theme['button_fg'])
        elif widget_type == 'Entry':
            widget.config(bg=theme['entry_bg'], fg=theme['entry_fg'],
                          insertbackground=theme['fg'])
        elif widget_type == 'Text':
            widget.config(bg=theme['entry_bg'], fg=theme['entry_fg'])
        elif widget_type in ['LabelFrame', 'Labelframe']:
            widget.config(bg=theme['label_frame_bg'], fg=theme['label_frame_fg'],
                          highlightbackground=theme['label_frame_bg'])
    except:
        pass
    
    # Рекурсивно применяем к дочерним элементам
    for child in widget.winfo_children():
        apply_theme_to_widget(child, theme)

def toggle_theme():
    """Переключает между темной и светлой темой"""
    global current_theme
    if current_theme == LIGHT_THEME:
        current_theme = DARK_THEME
        btn_theme.config(text="Светлая тема")
    else:
        current_theme = LIGHT_THEME
        btn_theme.config(text="Темная тема")
    apply_theme()

def set_window_size(size):
    """Устанавливает размер окна"""
    if size == "fullscreen":
        root.attributes('-fullscreen', True)
        btn_fullscreen.config(relief=tk.SUNKEN)
        btn_default.config(relief=tk.RAISED)
        btn_medium.config(relief=tk.RAISED)
        btn_large.config(relief=tk.RAISED)
        btn_little.config(relief=tk.RAISED)
    else:
        root.attributes('-fullscreen', False)
        if size == "default":
            root.geometry("1200x900")
            btn_default.config(relief=tk.SUNKEN)
            btn_fullscreen.config(relief=tk.RAISED)
            btn_medium.config(relief=tk.RAISED)
            btn_large.config(relief=tk.RAISED)
            btn_little.config(relief=tk.RAISED)
        elif size == "medium":
            root.geometry("1000x700")
            btn_medium.config(relief=tk.SUNKEN)
            btn_fullscreen.config(relief=tk.RAISED)
            btn_default.config(relief=tk.RAISED)
            btn_large.config(relief=tk.RAISED)
            btn_little.config(relief=tk.RAISED)
        elif size == "large":
            root.geometry("1400x1000")
            btn_large.config(relief=tk.SUNKEN)
            btn_fullscreen.config(relief=tk.RAISED)
            btn_default.config(relief=tk.RAISED)
            btn_medium.config(relief=tk.RAISED)
            btn_little.config(relief=tk.RAISED)
        elif size == "little":
            root.geometry("700x600")
            btn_little.config(relief=tk.SUNKEN)
            btn_large.config(relief=tk.RAISED)
            btn_fullscreen.config(relief=tk.RAISED)
            btn_default.config(relief=tk.RAISED)
            btn_medium.config(relief=tk.RAISED)

# === Функция для создания выписки ===
def generate_transaction_report():
    selected_client = table_clients_transactions.focus()
    if not selected_client:
        messagebox.showwarning("Ошибка", "Выберите клиента!")
        return
    
    client_id, client_name = table_clients_transactions.item(selected_client, "values")
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("SELECT passport FROM clients WHERE id = ?", (client_id,))
    passport = cursor.fetchone()[0]
    
    # Проверяем, выбран ли конкретный счет
    selected_account = tree_accounts_transactions.focus()
    single_account_mode = bool(selected_account)
    
    if single_account_mode:
        account_id, account_name, currency, balance = tree_accounts_transactions.item(selected_account, "values")
        
        # Получаем транзакции для выбранного счета
        cursor.execute("""
            SELECT t.date, t.type, t.amount, t.description 
            FROM transactions t
            WHERE t.account_id = ?
            ORDER BY t.date DESC
        """, (account_id,))
        transactions = cursor.fetchall()
        
        filename = f"{client_name}({passport}) {account_name} ({datetime.now().strftime('%Y-%m-%d %H%M%S')}).docx"
        save_folder = "Reports"
        os.makedirs(save_folder, exist_ok=True)
        full_path = os.path.join(save_folder, filename)
    
    else:
        if not messagebox.askyesno("Подтверждение", 
                                 "Счет не выбран. Создать выписку по ВСЕМ счетам клиента?"):
            conn.close()
            return
        
        # Получаем все счета клиента
        cursor.execute("SELECT id, account_name, balance FROM accounts WHERE client_id = ?", (client_id,))
        accounts = cursor.fetchall()
        
        if not accounts:
            messagebox.showwarning("Ошибка", "У клиента нет счетов!")
            conn.close()
            return
        
        # Получаем все транзакции по всем счетам клиента
        account_ids = [acc[0] for acc in accounts]
        placeholders = ','.join('?' for _ in account_ids)
        query = f"""
            SELECT t.date, a.account_name, t.type, t.amount, a.currency, t.description 
            FROM transactions t
            JOIN accounts a ON t.account_id = a.id
            WHERE t.account_id IN ({placeholders})
            ORDER BY t.date DESC
        """
        cursor.execute(query, account_ids)
        transactions = cursor.fetchall()
        
        filename = f"{client_name}({passport}) все счета ({datetime.now().strftime('%Y-%m-%d %H%M%S')}).docx"
        save_folder = "Reports"
        os.makedirs(save_folder, exist_ok=True)
        full_path = os.path.join(save_folder, filename)


    if not transactions:
        messagebox.showwarning("Ошибка", "Нет транзакций для выбранного режима!")
        conn.close()
        return
    
    doc = Document()
   
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title = doc.add_heading('Банковская выписка', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Клиент: {client_name} ({passport})', style='Intense Quote')
    
    # Информация о счете/счетах
    if single_account_mode:
        doc.add_paragraph(f'Счет: {account_name}', style='Intense Quote')
        cursor.execute("SELECT balance FROM accounts WHERE id = ?", (account_id,))
        balance = cursor.fetchone()[0]
        doc.add_paragraph(f'Текущий баланс: {balance:.2f} {currency}', style='Intense Quote')
    else:
        doc.add_paragraph('Все счета клиента', style='Intense Quote')
        
        accounts_table = doc.add_table(rows=1, cols=2)
        accounts_table.style = 'Table Grid'
        hdr_cells = accounts_table.rows[0].cells
        hdr_cells[0].text = 'Название'
        hdr_cells[1].text = 'Баланс'
        
        for acc in accounts:
            row_cells = accounts_table.add_row().cells
            row_cells[0].text = acc[1]
            row_cells[1].text = f"{acc[2]:.2f}"
    
    # Дата создания выписки
    current_date = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    doc.add_paragraph(f'Дата формирования: {current_date}')
    doc.add_paragraph()  
    
    # Добавляем таблицу с транзакциями
    trans_title = doc.add_heading('Список транзакций', level=2)
    if single_account_mode:
        table = doc.add_table(rows=1, cols=5)
    else:
         table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дата'
    hdr_cells[1].text = 'Счет' if not single_account_mode else 'Тип операции'
    hdr_cells[2].text = 'Тип операции' if not single_account_mode else 'Сумма'
    hdr_cells[3].text = 'Сумма' if not single_account_mode else 'Валюта'
    hdr_cells[4].text = 'Валюта' if not single_account_mode else 'Описание'
    
    if not single_account_mode:
        hdr_cells[5].text = 'Описание'
    
    for trans in transactions:
        if single_account_mode:
            date, trans_type, amount, description = trans
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = trans_type
            row_cells[2].text = f"{amount:.2f}"
            row_cells[3].text = currency
            row_cells[4].text = description if description else "-"
        else:
            date, account_name, trans_type, amount, currency, description = trans
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = account_name
            row_cells[2].text = trans_type
            row_cells[3].text = f"{amount:.2f}"
            row_cells[4].text = currency
            row_cells[5].text = description if description else "-"
    
    doc.add_paragraph("\n\n")
    doc.add_paragraph("Все комментарии и пожелании оставляйте старосте группы Ит-13", style='Intense Quote')
    
    try:
        doc.save(full_path)
        messagebox.showinfo("Успех", f"Выписка сохранена в файл:\n{full_path}")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{str(e)}")
    finally:
        conn.close()

# === Функции для работы с валютами ===
def get_currency_data():
    """Получает и возвращает данные о валютах с сайта ЦБ РФ"""
    try:
        response = requests.get('https://www.cbr-xml-daily.ru/daily_json.js')
        data = response.json()
        return {
            'USD': data['Valute']['USD']['Value'],
            'EUR': data['Valute']['EUR']['Value'],
            'RUB': 1.0
        }
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось получить курсы валют: {e}")
        return None

def get_currency_rates():
    """Возвращает форматированную таблицу курсов валют"""
    rates = get_currency_data()
    if not rates:
        return "Ошибка загрузки курсов валют"
    
    usd = rates['USD']
    eur = rates['EUR']
    rub = 1.0

    rub_to_usd = 1 / usd
    rub_to_eur = 1 / eur
    usd_to_rub = usd
    usd_to_eur = usd / eur
    eur_to_rub = eur
    eur_to_usd = eur / usd

    table = (
        f"{'Валюта':<10}{'RUB ':>10}{'USD ':>10}{'EUR ':>10}\n"
        f" "
        f"{'-'*51}\n"
        f"{'1 RUB':<10}{rub:>10.4f}{rub_to_usd:>10.4f}{rub_to_eur:>10.4f}\n"
        f"{' 1 USD':<10}{usd_to_rub:>10.4f}{1.0:>10.4f}{usd_to_eur:>10.4f}\n"
        f"{' 1 EUR':<10}{eur_to_rub:>10.4f}{eur_to_usd:>10.4f}{1.0:>10.4f}"
    )
    return table

def convert_currency(amount, from_currency, to_currency):
    if from_currency == to_currency:
        return amount
    
    rates = get_currency_data()
    if not rates:
        return None
    
    # Все конвертируем через RUB
    if from_currency != "RUB":
        amount = amount * rates[from_currency]
    
    if to_currency != "RUB":
        amount = amount / rates[to_currency]
        
    return amount

def update_currency_label():
    rates_text = get_currency_rates()
    lbl_money_rate.config(text=rates_text)

# === Инициализация базы данных ===
def init_db():
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()

    cursor.execute("""CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY, 
        name TEXT, 
        passport TEXT NOT NULL UNIQUE, 
        address TEXT)""")

    cursor.execute("""CREATE TABLE IF NOT EXISTS accounts (
        id INTEGER PRIMARY KEY, 
        client_id INTEGER, 
        account_name TEXT,
        balance REAL, 
        currency TEXT,
        FOREIGN KEY(client_id) REFERENCES clients(id))""")

    cursor.execute("""CREATE TABLE IF NOT EXISTS transactions (
    id INTEGER PRIMARY KEY, 
    account_id INTEGER, 
    type TEXT, 
    amount REAL, 
    date TEXT,
    description TEXT,
    related_account_id INTEGER,
    FOREIGN KEY(account_id) REFERENCES accounts(id))""")

    conn.commit()
    conn.close()

# === Общие функции ===

def get_client_name(client_id):
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM clients WHERE id = ?", (client_id,))
    result = cursor.fetchone()
    conn.close()
    return result[0] if result else "Неизвестный клиент"

def get_account_info(account_id):
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("""SELECT a.balance, a.currency, c.name, a.account_name 
                      FROM accounts a 
                      JOIN clients c ON a.client_id = c.id 
                      WHERE a.id = ?""", (account_id,))
    result = cursor.fetchone()
    conn.close()
    return result if result else (0, "", "Неизвестный счет", "")

# === Функции для сортировки таблиц ===
def treeview_sort_column(tv, col, reverse):
    l = [(tv.set(k, col), k) for k in tv.get_children('')]
    
    # Пытаемся преобразовать к числу, если возможно
    try:
        if col in ["Баланс", "Сумма", "ID"]:
            l.sort(key=lambda t: float(t[0]), reverse=reverse)
        elif col == "Дата":
            l.sort(key=lambda t: datetime.strptime(t[0], "%Y-%m-%d %H:%M:%S"), reverse=reverse)
        else:
            l.sort(key=lambda t: t[0], reverse=reverse)
    except ValueError:
        l.sort(key=lambda t: t[0], reverse=reverse)
    
    # Перемещаем элементы в отсортированном порядке
    for index, (val, k) in enumerate(l):
        tv.move(k, '', index)
    
    # Устанавливаем обратную сортировку для следующего клика
    tv.heading(col, command=lambda: treeview_sort_column(tv, col, not reverse))

############################################################################################################################################

# === Функции для работы с клиентами ===
def validate_passport(passport):
    return bool(re.fullmatch(r'[A-Z]{2}\d{7}', passport))

 # Приводим к верхнему регистру
def on_passport_entry(event):
    value = entry_passport.get()
    value = ''.join(c for c in value if c.isalnum())
    entry_passport.delete(0, tk.END)
    entry_passport.insert(0, value.upper())

def add_client():
    name = entry_name.get()
    passport = entry_passport.get()
    address = entry_address.get()

    if not name.strip() or not passport.strip() or not address.strip():
        messagebox.showwarning("Ошибка", "Заполните все поля!")
        return
    
    if not validate_passport(passport):
        messagebox.showwarning("Ошибка", "Паспорт должен быть в формате 2 латинские буквы и 7 цифр (например, AB1234567)")
        return

    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()

    try:
        cursor.execute("INSERT INTO clients (name, passport, address) VALUES (?, ?, ?)",
                       (name, passport, address))
        conn.commit()
        update_table_clients()
        clear_form()
    except sqlite3.IntegrityError:
        messagebox.showerror("Ошибка", "Клиент с таким паспортом уже существует!")
    finally:
        conn.close()
        entry_name.focus()

def search_clients():
    search_term = entry_search_client.get().strip().lower()

    if not search_term:
        query = "SELECT id, name, passport, address FROM clients"
        params = ()
    else:
        query = """
            SELECT id, name, passport, address 
            FROM clients 
            WHERE LOWER(name) LIKE ?
        """
        params = (f"%{search_term}%",)

    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()

    # Обновляем таблицу
    table_clients.delete(*table_clients.get_children())
    for row in rows:
        table_clients.insert("", tk.END, values=row)


def update_table_clients():
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, passport, address FROM clients")
    rows = cursor.fetchall()
    conn.close()
    
    # Очистка и обновление основной таблицы клиентов
    table_clients.delete(*table_clients.get_children())
    for row in rows:
        table_clients.insert("", tk.END, values=row)
    
    # Очистка и обновление мини-таблицы клиентов (на вкладке "Счета")
    table_clients_accounts.delete(*table_clients_accounts.get_children())
    for row in rows:
        table_clients_accounts.insert("", tk.END, values=(row[0], row[1]))  

    # Очистка и обновление мини-таблицы клиентов (на вкладке "Транзакции")
    table_clients_transactions.delete(*table_clients_transactions.get_children())
    for row in rows:
        table_clients_transactions.insert("", tk.END, values=(row[0], row[1]))
    
    entry_name.focus()  # Фокус на поле ввода имени

def edit_client():
    global selected_client_id, editing_mode
        
    selected_item = table_clients.focus()
    if not selected_item:
        messagebox.showwarning("Ошибка", "Выберите клиента для редактирования!")
        return

    values = table_clients.item(selected_item, "values")
    selected_client_id = values[0]
    editing_mode = True
    update_client_buttons()
    
    table_clients.config(selectmode="none")
    
    entry_name.delete(0, tk.END)
    entry_name.insert(0, values[1])
    entry_passport.delete(0, tk.END)
    entry_passport.insert(0, values[2])
    entry_address.delete(0, tk.END)
    entry_address.insert(0, values[3])

def save_client():
    global selected_client_id, editing_mode
        
    new_name = entry_name.get()
    new_passport = entry_passport.get()
    new_address = entry_address.get()

    if not new_name or not new_passport or not new_address:
        messagebox.showwarning("Ошибка", "Заполните все поля!")
        return

    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE clients SET name = ?, passport = ?, address = ? WHERE id = ?",
               (new_name, new_passport, new_address, selected_client_id))
        conn.commit()
    except sqlite3.IntegrityError:
        messagebox.showerror("Ошибка", "Клиент с таким паспортом уже существует!")
        return
    conn.commit()
    conn.close()
    update_table_clients()
    clear_form()
    table_clients.config(selectmode="browse")
    editing_mode = False
    selected_client_id = None
    update_client_buttons()

    messagebox.showinfo("Успех", "Изменения сохранены!")

def cancel_edit():
    global selected_client_id, editing_mode
    clear_form()
    table_clients.config(selectmode="browse")
    editing_mode = False
    selected_client_id = None
    update_client_buttons()         

def clear_form():
    entry_name.delete(0, tk.END)
    entry_passport.delete(0, tk.END)
    entry_address.delete(0, tk.END)

def delete_client():
    if editing_mode:
        messagebox.showwarning("Ошибка", "Завершите текущее редактирование!")
        return
        
    selected_item = table_clients.focus()
    if not selected_item:
        messagebox.showwarning("Ошибка", "Выберите клиента для удаления!")
        return

    values = table_clients.item(selected_item, "values")
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM accounts WHERE client_id = ?", (values[0],))
    account_count = cursor.fetchone()[0]
    
    if account_count > 0:
        messagebox.showwarning("Ошибка", "Нельзя удалить клиента с существующими счетами!")
        conn.close()
        return
    
    if messagebox.askyesno("Подтверждение", f"Удалить клиента {values[1]}?"):
        cursor.execute("DELETE FROM clients WHERE id = ?", (values[0],))
        conn.commit()
        conn.close()
        update_table_clients()


def update_client_buttons():
    for widget in button_frame.winfo_children():
        widget.pack_forget()

    if editing_mode:
        btn_add_client.pack_forget()
        btn_delete_client.pack_forget()
        btn_edit_client.pack_forget()
        btn_clear.pack(side=tk.LEFT, padx=5)
        btn_save.pack(side=tk.LEFT, padx=5)
        btn_cancel.pack(side=tk.LEFT, padx=5)
    else:
        btn_add_client.pack(side=tk.LEFT, padx=5)
        btn_delete_client.pack(side=tk.LEFT, padx=5)
        btn_edit_client.pack(side=tk.LEFT, padx=5)
        btn_clear.pack(side=tk.LEFT, padx=5)
        btn_save.pack_forget()
        btn_cancel.pack_forget()

############################################################################################################################################

# === Функции для работы со счетами ===
def search_accounts():
    search_term = entry_search_account.get().strip()
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    if search_term:
        cursor.execute("""SELECT a.id, c.name || ' (ID=' || c.id || ')', a.account_name, a.currency, ROUND(a.balance, 2)  
                          FROM accounts a 
                          JOIN clients c ON a.client_id = c.id
                          WHERE a.account_name COLLATE NOCASE LIKE ?""", 
                      (f"%{search_term}%",))
    else:
        cursor.execute("""SELECT a.id, c.name || ' (ID=' || c.id || ')', a.account_name, a.currency, ROUND(a.balance, 2)  
                          FROM accounts a 
                          JOIN clients c ON a.client_id = c.id""")
    
    rows = cursor.fetchall()
    conn.close()
    
    table_accounts.delete(*table_accounts.get_children())
    for row in rows:
        table_accounts.insert("", tk.END, values=row)


def add_account():
    selected_client = table_clients_accounts.focus()
    if not selected_client:
        messagebox.showwarning("Ошибка", "Выберите клиента!")
        return
        
    client_id = table_clients_accounts.item(selected_client, "values")[0]
    account_name = entry_account_name.get().strip()
    currency = combo_currency.get()
    initial_balance = entry_initial_balance.get()

    if not currency or not initial_balance or not account_name:
        messagebox.showwarning("Ошибка", "Заполните все поля!")
        return
        
    try:
        initial_balance = float(initial_balance)
    except ValueError:
        messagebox.showwarning("Ошибка", "Введите корректную сумму!")
        return

    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    # Проверяем, есть ли уже счет с таким названием у этого клиента
    cursor.execute("SELECT COUNT(*) FROM accounts WHERE client_id = ? AND account_name = ?", 
                   (client_id, account_name))
    existing_accounts = cursor.fetchone()[0]
    
    if existing_accounts > 0:
        messagebox.showwarning("Ошибка", "У этого клиента уже есть счет с таким названием!")
        conn.close()
        return
    
    cursor.execute("INSERT INTO accounts (client_id, account_name, balance, currency) VALUES (?, ?, ?, ?)", 
                   (client_id, account_name, initial_balance, currency))
    
    account_id = cursor.lastrowid
    
    cursor.execute("""INSERT INTO transactions (account_id, type, amount, date, description) 
                      VALUES (?, ?, ?, ?, ?)""",
                   (account_id, "Создание", initial_balance, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                    f"Открытие счета '{account_name}'"))
    
    conn.commit()
    conn.close()
    
    update_accounts_tree()
    update_transactions_tree()
    entry_initial_balance.delete(0, tk.END)
    entry_account_name.delete(0, tk.END)

def update_accounts_tree():
    # Обновляем основную таблицу счетов (все счета)
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    cursor.execute("""SELECT a.id, c.name || ' (ID=' || c.id || ')', a.account_name, a.currency, ROUND(a.balance, 2)  
                      FROM accounts a 
                      JOIN clients c ON a.client_id = c.id""")
    all_accounts = cursor.fetchall()
    conn.close()

    table_accounts.delete(*table_accounts.get_children())
    for account in all_accounts:
        table_accounts.insert("", tk.END, values=account)

    # Обновляем мини-таблицу счетов выбранного клиента
    selected_client = table_clients_accounts.focus()
    if selected_client:
        client_id = table_clients_accounts.item(selected_client, "values")[0]
        client_name = table_clients_accounts.item(selected_client, "values")[1]
        lbl_accounts_client.config(text=f"Клиент: {client_name}")
        
        conn = sqlite3.connect("bank.db")
        cursor = conn.cursor()
        cursor.execute("""SELECT a.id, a.account_name, a.currency, ROUND(a.balance, 2)  
                          FROM accounts a 
                          WHERE a.client_id = ?""", (client_id,))
        client_accounts = cursor.fetchall()
        conn.close()

        tree_accounts_list.delete(*tree_accounts_list.get_children())
        for account in client_accounts:
            tree_accounts_list.insert("", tk.END, values=account)


def delete_account():
    selected_account = table_accounts.focus()
    if not selected_account:
        messagebox.showwarning("Ошибка", "Выберите счет для удаления!")
        return

    account_id = table_accounts.item(selected_account, "values")[0]
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    cursor.execute("SELECT balance FROM accounts WHERE id = ?", (account_id,))
    balance = cursor.fetchone()[0]
    
    if balance != 0:
        messagebox.showwarning("Ошибка", "Нельзя удалить счет с ненулевым балансом!")
        conn.close()
        return
    
    if messagebox.askyesno("Подтверждение", "Удалить выбранный счет?"):
        cursor.execute("DELETE FROM transactions WHERE account_id = ?", (account_id,))
        cursor.execute("DELETE FROM accounts WHERE id = ?", (account_id,))
        conn.commit()
        conn.close()
        update_accounts_tree()
        update_transactions_tree()
        clear_transactions_table()

############################################################################################################################################

# === Функции для работы с транзакциями ===
def update_transfer_info(event=None):
    """Обновляет информацию о переводе (комиссия, итоговая сумма)"""
    if combo_transaction_type.get() != "Перевод":
        lbl_commission.config(text="")
        lbl_final_amount.config(text="")
        return
        
    try:
        amount = float(entry_transaction_amount.get())
    except:
        lbl_commission.config(text="Комиссия: -")
        lbl_final_amount.config(text="Итого: -")
        return
    
    selected_account = tree_accounts_transactions.focus()
    selected_transfer = tree_transfer_to.focus()
    
    if not selected_account or not selected_transfer:
        lbl_commission.config(text="Комиссия: -")
        lbl_final_amount.config(text="Итого: -")
        return
    
    # Получаем информацию о счетах
    from_account_id = tree_accounts_transactions.item(selected_account, "values")[0]
    to_account_id = tree_transfer_to.item(selected_transfer, "values")[0]
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    # Получаем валюту и клиента для исходного счета
    cursor.execute("SELECT currency, client_id FROM accounts WHERE id = ?", (from_account_id,))
    from_currency, from_client_id = cursor.fetchone()
    
    # Получаем валюту и клиента для счета получателя
    cursor.execute("SELECT currency, client_id FROM accounts WHERE id = ?", (to_account_id,))
    to_currency, to_client_id = cursor.fetchone()
    
    conn.close()
    
    # Конвертируем сумму
    converted_amount = convert_currency(amount, from_currency, to_currency)
    if converted_amount is None:
        return
    
    # Проверяем, нужно ли брать комиссию
    if from_client_id == to_client_id:
        commission = 0
        commission_text = "Комиссия: 0% (перевод между своими счетами)"
    else:
        commission = 0.1  
        commission_text = f"Комиссия: 10% ({amount * commission:.2f} {from_currency})"
    
    final_amount = amount * (1 - commission)
    final_converted = convert_currency(final_amount, from_currency, to_currency)
    
    lbl_commission.config(text=commission_text)
    lbl_final_amount.config(text=f"Итого будет переведено: {final_converted:.2f} {to_currency}")

def on_transaction_type_change(event):
    if combo_transaction_type.get() == "Перевод":
        selected_account = tree_accounts_transactions.focus()
        if selected_account:
            account_id = tree_accounts_transactions.item(selected_account, "values")[0]
            update_transfer_accounts(account_id)
        transfer_frame.pack()
        lbl_final_amount.pack()
        lbl_commission.pack()
        update_transfer_info()
    else:
        transfer_frame.pack_forget()
        lbl_final_amount.pack_forget()
        lbl_commission.pack_forget()

def update_transfer_accounts(account_id):
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    # Получаем ID клиента текущего счета
    cursor.execute("SELECT client_id FROM accounts WHERE id = ?", (account_id,))
    current_client_id = cursor.fetchone()[0]
    
    # Получаем все счета, кроме текущего
    cursor.execute("""SELECT a.id, a.account_name, a.currency, ROUND(a.balance, 2), c.name, c.id
                      FROM accounts a 
                      JOIN clients c ON a.client_id = c.id
                      WHERE a.id != ?""", (account_id,))
    accounts = cursor.fetchall()
    conn.close()
    
    tree_transfer_to.delete(*tree_transfer_to.get_children())
    for account in accounts:
        tree_transfer_to.insert("", tk.END, values=account)
    


def update_transactions_tree():
    selected_client = table_clients_transactions.focus()
    if selected_client:
        client_id = table_clients_transactions.item(selected_client, "values")[0]
        client_name = table_clients_transactions.item(selected_client, "values")[1]
        lbl_transactions_client.config(text=f"Клиент: {client_name}")
        lbl_transactions_client_account.config(text="Выберите счет")
        conn = sqlite3.connect("bank.db")
        cursor = conn.cursor()
        cursor.execute("""SELECT a.id, a.account_name, a.currency, ROUND(a.balance, 2)  
                          FROM accounts a 
                          WHERE a.client_id = ?""", (client_id,))
        client_accounts = cursor.fetchall()

        conn.close()

        tree_accounts_transactions.delete(*tree_accounts_transactions.get_children())
        for account in client_accounts:
            tree_accounts_transactions.insert("", tk.END, values=account)

def clear_transactions_table():
    table_transactions.delete(*table_transactions.get_children())
    lbl_transactions_account.config(text="Транзакции для (счет не выбран)")

def show_transactions(account_id=None):
    if account_id is None:
        selected_account = tree_accounts_transactions.focus()
        if not selected_account:
            return
        account_id = tree_accounts_transactions.item(selected_account, "values")[0]
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    # Получаем информацию о счете для заголовка
    cursor.execute("""SELECT a.account_name, a.currency, c.name, c.id 
                      FROM accounts a 
                      JOIN clients c ON a.client_id = c.id 
                      WHERE a.id = ?""", (account_id,))
    account_info = cursor.fetchone()
    
    if account_info:
        account_name, currency, client_name, client_id = account_info
        lbl_transactions_account.config(
            text=f"Транзакции для счета: {account_name} (Валюта: {currency}, Владелец: {client_name} (ID={client_id}))"
        )
        selected_account = tree_accounts_transactions.focus()
        if selected_account:
            lbl_transactions_client_account.config(text=f"Счет: {account_name}")
        else:
            lbl_transactions_client_account.config(text=f"Выберите счет")

    cursor.execute("""SELECT t.id, t.type, t.amount, t.date, t.description 
                      FROM transactions t 
                      WHERE t.account_id = ? 
                      ORDER BY t.date DESC""", (account_id,))
    transactions = cursor.fetchall()
    conn.close()
    
    table_transactions.delete(*table_transactions.get_children())
    for transaction in transactions:
        table_transactions.insert("", tk.END, values=transaction)

def add_transaction():
    selected_account = tree_accounts_transactions.focus()
    if not selected_account:
        messagebox.showwarning("Ошибка", "Выберите счет!")
        return
    
    account_id = tree_accounts_transactions.item(selected_account, "values")[0]
    transaction_type = combo_transaction_type.get()
    amount = entry_transaction_amount.get()
    description = entry_transaction_desc.get().strip()
    
    if not transaction_type or not amount:
        messagebox.showwarning("Ошибка", "Заполните все поля!")
        return
    
    try:
        amount = float(amount)
        if amount <= 0:
            messagebox.showwarning("Ошибка", "Сумма должна быть положительной!")
            return
    except ValueError:
        messagebox.showwarning("Ошибка", "Введите корректную сумму!")
        return
    
    conn = sqlite3.connect("bank.db")
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT ROUND(balance, 2), currency, client_id FROM accounts WHERE id = ?", (account_id,))
        balance, from_currency, from_client_id = cursor.fetchone()
        
        if transaction_type in ["Снятие", "Перевод"] and balance < amount:
            messagebox.showwarning("Ошибка", "Недостаточно средств на счете!")
            return
            
        if transaction_type == "Перевод":
            selected_transfer = tree_transfer_to.focus()
            if not selected_transfer:
                messagebox.showwarning("Ошибка", "Выберите счет для перевода!")
                return
            to_account_id = tree_transfer_to.item(selected_transfer, "values")[0]
            
            # Получаем информацию о счете получателя
            cursor.execute("""SELECT a.currency, a.client_id, a.account_name, c.name 
                             FROM accounts a 
                             JOIN clients c ON a.client_id = c.id 
                             WHERE a.id = ?""", (to_account_id,))
            to_currency, to_client_id, to_account_name, to_client_name = cursor.fetchone()
            
            # Конвертируем сумму
            converted_amount = convert_currency(amount, from_currency, to_currency)
            if converted_amount is None:
                return
            
            # Определяем комиссию
            commission = 0.1 if from_client_id != to_client_id else 0
            final_amount = amount * (1 - commission)
            final_converted = convert_currency(final_amount, from_currency, to_currency)
            
            # Обновляем баланс исходного счета (списываем полную сумму)
            cursor.execute("UPDATE accounts SET balance = balance - ? WHERE id = ?", (amount, account_id))
            
            # Формируем описание для транзакции
            full_description = f'Перевод на счет "{to_account_name}" пользователя {to_client_name}'
            if from_client_id == to_client_id:
                full_description += " (между своими счетами)"
            else:
                full_description += f" (комиссия 10%: {amount * commission:.2f} {from_currency})"
            
            if description:
                full_description += f" ({description})"
            
            # Добавляем транзакцию списания
            cursor.execute("""INSERT INTO transactions 
                              (account_id, type, amount, date, description, related_account_id) 
                              VALUES (?, ?, ?, ?, ?, ?)""",
                           (account_id, "Перевод", amount, 
                            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                            full_description, to_account_id))
            
            # Зачисляем на другой счет (уже с учетом комиссии и конвертации)
            cursor.execute("UPDATE accounts SET balance = balance + ? WHERE id = ?", 
                          (final_converted, to_account_id))
            
            # Формируем описание для получателя
            cursor.execute("""SELECT a.account_name, c.name 
                             FROM accounts a 
                             JOIN clients c ON a.client_id = c.id 
                             WHERE a.id = ?""", (account_id,))
            from_account_name, from_client_name = cursor.fetchone()
            
            recipient_description = f'Перевод со счета "{from_account_name}" пользователя {from_client_name}'
            if description:
                recipient_description += f" ({description})"
            
            # Добавляем транзакцию зачисления
            cursor.execute("""INSERT INTO transactions 
                              (account_id, type, amount, date, description, related_account_id) 
                              VALUES (?, ?, ?, ?, ?, ?)""",
                           (to_account_id, "Пополнение", final_converted, 
                            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                            recipient_description, account_id))
        
        else:  # Пополнение или снятие
            new_balance = balance - amount if transaction_type == "Снятие" else balance + amount
            cursor.execute("UPDATE accounts SET balance = ? WHERE id = ?", (new_balance, account_id))
            
            # Добавляем транзакцию
            cursor.execute("""INSERT INTO transactions 
                              (account_id, type, amount, date, description) 
                              VALUES (?, ?, ?, ?, ?)""",
                           (account_id, transaction_type, amount, 
                            datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 
                            description if description else ""))
        
        conn.commit()
        
        # Обновляем отображение
        update_transactions_tree()
        update_accounts_tree()
        show_transactions(account_id)
        update_transfer_accounts(account_id)
        
        # Очищаем поля ввода
        entry_transaction_amount.delete(0, tk.END)
        entry_transaction_desc.delete(0, tk.END)
        
        messagebox.showinfo("Успех", "Транзакция выполнена!")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")
    finally:
        conn.close()

############################################################################################################################################

# === Графический интерфейс ===
root = tk.Tk()
root.title("Банковская система")
root.iconphoto(True, tk.PhotoImage(file='icon.png'))
root.geometry("1200x900")


# Глобальные переменные для управления состоянием редактирования
selected_client_id = None
editing_mode = False

notebook = ttk.Notebook(root)
frame_clients = ttk.Frame(notebook)
frame_accounts = ttk.Frame(notebook)
frame_transactions = ttk.Frame(notebook)
frame_settings = ttk.Frame(notebook)


notebook.add(frame_clients, text="Клиенты")
notebook.add(frame_accounts, text="Счета")
notebook.add(frame_transactions, text="Транзакции")
notebook.add(frame_settings, text="Настройки")
notebook.pack(expand=True, fill="both")



# === Вкладка клиентов ===
client_form_frame = tk.Frame(frame_clients)
client_form_frame.pack(pady=10)

tk.Label(client_form_frame, text="ФИО:").grid(row=0, column=0, sticky="e")
entry_name = tk.Entry(client_form_frame, width=40)
entry_name.grid(row=0, column=1, padx=5, pady=5)

tk.Label(client_form_frame, text="Паспорт:").grid(row=1, column=0, sticky="e")
entry_passport = tk.Entry(client_form_frame, width=40)
entry_passport.grid(row=1, column=1, padx=5, pady=5)

tk.Label(client_form_frame, text="Адрес:").grid(row=2, column=0, sticky="e")
entry_address = tk.Entry(client_form_frame, width=40)
entry_address.grid(row=2, column=1, padx=5, pady=5)


button_frame = tk.Frame(frame_clients)
button_frame.pack(pady=5)
btn_add_client = tk.Button(button_frame, text="Добавить клиента", command=add_client)
btn_edit_client = tk.Button(button_frame, text="Редактировать клиента", command=edit_client)
btn_delete_client = tk.Button(button_frame, text="Удалить клиента", command=delete_client)
btn_clear = tk.Button(button_frame, text="Очистить форму", command=clear_form)
btn_add_client.pack(side=tk.LEFT, padx=5)
btn_edit_client.pack(side=tk.LEFT, padx=5)
btn_delete_client.pack(side=tk.LEFT, padx=5)
btn_clear.pack(side=tk.LEFT, padx=5)
btn_save = tk.Button(button_frame, text="Сохранить", command=save_client)
btn_cancel = tk.Button(button_frame, text="Отмена", command=cancel_edit)

# Поле поиска
search_client_frame = tk.Frame(frame_clients)
search_client_frame.pack(fill=tk.X, padx=10, pady=5)
tk.Label(search_client_frame, text="Поиск по ФИО:").pack(side=tk.LEFT)
entry_search_client = tk.Entry(search_client_frame, width=30)
entry_search_client.pack(side=tk.LEFT, padx=5)

btn_clear_search_client = tk.Button(search_client_frame, text="Сброс", command=lambda: [entry_search_client.delete(0, tk.END), update_table_clients()], width=7)
btn_clear_search_client.pack(side=tk.LEFT)
entry_search_client.bind("<KeyRelease>", lambda e: search_clients())

# Таблица клиентов
table_clients = ttk.Treeview(frame_clients, columns=("ID", "ФИО", "Паспорт", "Адрес"), show="headings", height=10)
table_clients.heading("ID", text="ID", command=lambda: treeview_sort_column(table_clients, "ID", False))
table_clients.heading("ФИО", text="ФИО", command=lambda: treeview_sort_column(table_clients, "ФИО", False))
table_clients.heading("Паспорт", text="Паспорт", command=lambda: treeview_sort_column(table_clients, "Паспорт", False))
table_clients.heading("Адрес", text="Адрес", command=lambda: treeview_sort_column(table_clients, "Адрес", False))
table_clients.column("ID", width=50)
table_clients.column("ФИО", width=200)
table_clients.column("Паспорт", width=150)
table_clients.column("Адрес", width=250)
table_clients.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

entry_passport.bind("<KeyRelease>", on_passport_entry)

############################################################################################################################################

# === Вкладка счетов ===
accounts_top_frame = tk.Frame(frame_accounts)
accounts_top_frame.pack(fill=tk.BOTH, expand=True)

# Разделение фрейма на 2 колонки
accounts_top_frame.columnconfigure(0, weight=1)  
accounts_top_frame.columnconfigure(1, weight=1) 

# Таблица клиентов (миниатюрная)
clients_frame_accounts = tk.Frame(accounts_top_frame)
clients_frame_accounts.grid(row=0, column=0, padx=10, pady=5, sticky="nsew")

tk.Label(clients_frame_accounts, text="Клиенты").pack()
table_clients_accounts = ttk.Treeview(clients_frame_accounts, columns=("ID", "ФИО"), show="headings", height=8)
table_clients_accounts.heading("ID", text="ID", command=lambda: treeview_sort_column(table_clients_accounts, "ID", False))
table_clients_accounts.heading("ФИО", text="ФИО", command=lambda: treeview_sort_column(table_clients_accounts, "ФИО", False))
table_clients_accounts.column("ID", width=50)
table_clients_accounts.column("ФИО", width=150)
table_clients_accounts.pack(expand=True, fill="both")

# Таблица счетов клиента (миниатюрная)
accounts_list_frame = tk.Frame(accounts_top_frame)
accounts_list_frame.grid(row=0, column=1, padx=10, pady=5, sticky="nsew")

tk.Label(accounts_list_frame, text="Счета клиента").pack()
tree_accounts_list = ttk.Treeview(accounts_list_frame, columns=("ID", "Название", "Валюта", "Баланс"), show="headings", height=8, selectmode="none")
tree_accounts_list.heading("ID", text="ID", command=lambda: treeview_sort_column(tree_accounts_list, "ID", False))
tree_accounts_list.heading("Название", text="Название", command=lambda: treeview_sort_column(tree_accounts_list, "Название", False))
tree_accounts_list.heading("Валюта", text="Валюта", command=lambda: treeview_sort_column(tree_accounts_list, "Валюта", False))
tree_accounts_list.heading("Баланс", text="Баланс", command=lambda: treeview_sort_column(tree_accounts_list, "Баланс", False))
tree_accounts_list.column("ID", width=50)
tree_accounts_list.column("Название", width=150)
tree_accounts_list.column("Валюта", width=80)
tree_accounts_list.column("Баланс", width=80)
tree_accounts_list.pack(expand=True, fill="both")

# Информация о выбранном клиенте
lbl_accounts_client = tk.Label(accounts_top_frame, text="Выберите клиента", font=('Arial', 10, 'bold'))
lbl_accounts_client.grid(row=1, column=0, padx=20, sticky="w")

# Форма добавления счетов
account_form_frame = tk.Frame(frame_accounts)
account_form_frame.pack(pady=10)

tk.Label(account_form_frame, text="Название счета:").grid(row=0, column=0, sticky="e")
entry_account_name = tk.Entry(account_form_frame, width=30)
entry_account_name.grid(row=0, column=1, padx=5, pady=5)

tk.Label(account_form_frame, text="Валюта:").grid(row=1, column=0, sticky="e")
combo_currency = ttk.Combobox(account_form_frame, values=["RUB", "USD", "EUR"], state="readonly", width=27)
combo_currency.grid(row=1, column=1, padx=5, pady=5)
combo_currency.current(0)

tk.Label(account_form_frame, text="Начальный баланс:").grid(row=2, column=0, sticky="e")
entry_initial_balance = tk.Entry(account_form_frame, width=30)
entry_initial_balance.grid(row=2, column=1, padx=5, pady=5)

btn_add_account = tk.Button(frame_accounts, text="Добавить счет", command=add_account, width=20)
btn_add_account.pack(pady=5)

btn_delete_account = tk.Button(frame_accounts, text="Удалить счет", command=delete_account, width=20)
btn_delete_account.pack(pady=5)


# Поле поиска
search_account_frame = tk.Frame(frame_accounts)
search_account_frame.pack(fill=tk.X, padx=10, pady=5)
tk.Label(search_account_frame, text="Поиск по названию счета:").pack(side=tk.LEFT)
entry_search_account = tk.Entry(search_account_frame, width=30)
entry_search_account.pack(side=tk.LEFT, padx=5)

btn_clear_search_account = tk.Button(search_account_frame, text="Сброс", command=lambda: [entry_search_account.delete(0, tk.END), update_accounts_tree()], width=10)
btn_clear_search_account.pack(side=tk.LEFT)
entry_search_account.bind("<KeyRelease>", lambda e: search_accounts())

# Основная таблица всех счетов
table_accounts = ttk.Treeview(frame_accounts, columns=("ID", "Владелец", "Название", "Валюта", "Баланс"), show="headings", height=20)
table_accounts.heading("ID", text="ID", command=lambda: treeview_sort_column(table_accounts, "ID", False))
table_accounts.heading("Владелец", text="Владелец", command=lambda: treeview_sort_column(table_accounts, "Владелец", False))
table_accounts.heading("Название", text="Название", command=lambda: treeview_sort_column(table_accounts, "Название", False))
table_accounts.heading("Валюта", text="Валюта", command=lambda: treeview_sort_column(table_accounts, "Валюта", False))
table_accounts.heading("Баланс", text="Баланс", command=lambda: treeview_sort_column(table_accounts, "Баланс", False))
table_accounts.column("ID", width=25)
table_accounts.column("Владелец", width=220)
table_accounts.column("Название", width=170)
table_accounts.column("Валюта", width=40)
table_accounts.column("Баланс", width=100)
table_accounts.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

# Привязка событий
table_clients_accounts.bind("<<TreeviewSelect>>", lambda event: update_accounts_tree())

############################################################################################################################################

# === Вкладка транзакций ===
transactions_top_frame = tk.Frame(frame_transactions)
transactions_top_frame.pack(fill=tk.BOTH, expand=True)

# Разделение фрейма на 2 колонки
transactions_top_frame.columnconfigure(0, weight=1)  
transactions_top_frame.columnconfigure(1, weight=1)  

# Таблица клиентов (миниатюрная)
clients_frame_transactions = tk.Frame(transactions_top_frame)
clients_frame_transactions.grid(row=0, column=0, padx=10, pady=5, sticky="nsew")

tk.Label(clients_frame_transactions, text="Клиенты").pack()
table_clients_transactions = ttk.Treeview(clients_frame_transactions, columns=("ID", "ФИО"), show="headings", height=8)
table_clients_transactions.heading("ID", text="ID", command=lambda: treeview_sort_column(table_clients_transactions, "ID", False))
table_clients_transactions.heading("ФИО", text="ФИО", command=lambda: treeview_sort_column(table_clients_transactions, "ФИО", False))
table_clients_transactions.column("ID", width=50)
table_clients_transactions.column("ФИО", width=150)
table_clients_transactions.pack(expand=True, fill="both")

# Таблица счетов клиента (миниатюрная)
accounts_transactions_frame = tk.Frame(transactions_top_frame)
accounts_transactions_frame.grid(row=0, column=1, padx=10, pady=5, sticky="nsew")

tk.Label(accounts_transactions_frame, text="Счета клиента").pack()
tree_accounts_transactions = ttk.Treeview(accounts_transactions_frame, columns=("ID", "Название", "Валюта", "Баланс"), show="headings", height=8)
tree_accounts_transactions.heading("ID", text="ID", command=lambda: treeview_sort_column(tree_accounts_transactions, "ID", False))
tree_accounts_transactions.heading("Название", text="Название", command=lambda: treeview_sort_column(tree_accounts_transactions, "Название", False))
tree_accounts_transactions.heading("Валюта", text="Валюта", command=lambda: treeview_sort_column(tree_accounts_transactions, "Валюта", False))
tree_accounts_transactions.heading("Баланс", text="Баланс", command=lambda: treeview_sort_column(tree_accounts_transactions, "Баланс", False))
tree_accounts_transactions.column("ID", width=50)
tree_accounts_transactions.column("Название", width=150)
tree_accounts_transactions.column("Валюта", width=80)
tree_accounts_transactions.column("Баланс", width=80)
tree_accounts_transactions.pack(expand=True, fill="both")

lbl_transactions_client = tk.Label(transactions_top_frame, text="Выберите клиента", font=('Arial', 10, 'bold'))
lbl_transactions_client.grid(row=1, column=0, padx=20, sticky="w")

lbl_transactions_client_account = tk.Label(transactions_top_frame, text="Выберите счет", font=('Arial', 10, 'bold'))
lbl_transactions_client_account.grid(row=1, column=1, padx=20, sticky="w")

# Форма добавления транзакций
transaction_form_frame = tk.Frame(frame_transactions)
transaction_form_frame.pack(pady=10)

tk.Label(transaction_form_frame, text="Тип операции:").grid(row=0, column=0, sticky="e")
combo_transaction_type = ttk.Combobox(transaction_form_frame, values=["Пополнение", "Снятие", "Перевод"], state="readonly", width=27)
combo_transaction_type.grid(row=0, column=1, padx=5, pady=5)
combo_transaction_type.current(0)
combo_transaction_type.bind("<<ComboboxSelected>>", on_transaction_type_change)

tk.Label(transaction_form_frame, text="Сумма:").grid(row=1, column=0, sticky="e")
entry_transaction_amount = tk.Entry(transaction_form_frame, width=30)
entry_transaction_amount.grid(row=1, column=1, padx=5, pady=5)

tk.Label(transaction_form_frame, text="Описание (необязательно):").grid(row=2, column=0, sticky="e")
entry_transaction_desc = tk.Entry(transaction_form_frame, width=30)
entry_transaction_desc.grid(row=2, column=1, padx=5, pady=5)

transaction_form1_frame = tk.Frame(frame_transactions)
transaction_form1_frame.pack()

lbl_commission = tk.Label(transaction_form1_frame, text="Комиссия: -", fg="red", font=('Arial', 10, 'bold'))
lbl_final_amount = tk.Label(transaction_form1_frame, text="Итого: -", fg="green", font=('Arial', 10, 'bold'))

# Кнопки
btn_add_transaction = tk.Button(frame_transactions, text="Выполнить операцию", command=add_transaction, width=20)
btn_add_transaction.pack(pady=5)

btn_generate_report = tk.Button(frame_transactions, text="Создать выписку", 
                              command=generate_transaction_report, width=20)
btn_generate_report.pack(pady=5)

# Фрейм для переводов и курсов валют
transfer_info_frame = tk.Frame(frame_transactions)
transfer_info_frame.pack(fill=tk.X, padx=10, pady=5)

# Внутренний фрейм — делим на левый и правый
transfer_frame = tk.Frame(transfer_info_frame)
transfer_frame.pack(fill=tk.X, pady=5)

# Левый
left_transfer_frame = tk.Frame(transfer_frame)
left_transfer_frame.pack(side="left", fill=tk.BOTH, expand=True, padx=(0, 10))

lbl_transfer_to = tk.Label(left_transfer_frame, text="Перевести на счет:", font=('Arial', 10, 'bold'))
lbl_transfer_to.pack(anchor="n")

tree_transfer_to = ttk.Treeview(left_transfer_frame, columns=("ID", "Название", "Валюта", "Баланс", "Владелец"), 
                               show="headings", height=5)
tree_transfer_to.heading("ID", text="ID", command=lambda: treeview_sort_column(tree_transfer_to, "ID", False))
tree_transfer_to.heading("Название", text="Название", command=lambda: treeview_sort_column(tree_transfer_to, "Название", False))
tree_transfer_to.heading("Валюта", text="Валюта", command=lambda: treeview_sort_column(tree_transfer_to, "Валюта", False))
tree_transfer_to.heading("Баланс", text="Баланс", command=lambda: treeview_sort_column(tree_transfer_to, "Баланс", False))
tree_transfer_to.heading("Владелец", text="Владелец", command=lambda: treeview_sort_column(tree_transfer_to, "Владелец", False))
tree_transfer_to.column("ID", width=30)
tree_transfer_to.column("Название", width=100)
tree_transfer_to.column("Валюта", width=50)
tree_transfer_to.column("Баланс", width=60)
tree_transfer_to.column("Владелец", width=150)
tree_transfer_to.pack(fill=tk.BOTH, expand=True)

# Правый
right_currency_frame = tk.Frame(transfer_frame)
right_currency_frame.pack(side="right", fill=tk.Y)

lbl_title = tk.Label(right_currency_frame, text="Курсы валют:", font=('Arial', 10, 'bold'))
lbl_title.pack(anchor="n", pady=(0, 5))

frame_border = tk.Frame(right_currency_frame, bd=2, relief="groove")  

lbl_money_rate = tk.Label(frame_border, text="Загрузка...", justify='right', font=('Arial', 9), anchor='e')
lbl_money_rate.pack(anchor="e", padx=5, pady=5)

frame_border.pack(anchor="e", pady=5, padx=5)

update_currency_label()

# Пустой фрейм для нормального скрытия пространства
account_info_frame = tk.Frame(transfer_info_frame)
account_info_frame.pack(fill=tk.X, pady=1)

# Скрываем по умолчанию
transfer_frame.pack_forget()

# Таблица транзакций
lbl_transactions_account = tk.Label(frame_transactions, text="Транзакции для (счет не выбран)", font=('Arial', 10, 'bold'))
lbl_transactions_account.pack()

table_transactions = ttk.Treeview(frame_transactions, columns=("ID", "Тип", "Сумма", "Дата", "Описание"), show="headings", height=20)
table_transactions.heading("ID", text="ID", command=lambda: treeview_sort_column(table_transactions, "ID", False))
table_transactions.heading("Тип", text="Тип", command=lambda: treeview_sort_column(table_transactions, "Тип", False))
table_transactions.heading("Сумма", text="Сумма", command=lambda: treeview_sort_column(table_transactions, "Сумма", False))
table_transactions.heading("Дата", text="Дата", command=lambda: treeview_sort_column(table_transactions, "Дата", False))
table_transactions.heading("Описание", text="Описание")
table_transactions.column("ID", width=50)
table_transactions.column("Тип", width=75)
table_transactions.column("Сумма", width=75)
table_transactions.column("Дата", width=100)
table_transactions.column("Описание", width=400)
table_transactions.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)


# Привязка событий
table_clients_transactions.bind("<<TreeviewSelect>>", lambda event: update_transactions_tree())
def on_account_select(event):
    selected_item = tree_accounts_transactions.focus()
    if selected_item:  
        account_id = tree_accounts_transactions.item(selected_item, "values")[0]
        show_transactions(account_id)
        if combo_transaction_type.get() == "Перевод":
            update_transfer_accounts(account_id)

tree_accounts_transactions.bind("<<TreeviewSelect>>", on_account_select)
entry_transaction_amount.bind("<KeyRelease>", update_transfer_info)
tree_transfer_to.bind("<<TreeviewSelect>>", lambda e: update_transfer_info())

############################################################################################################################################

# Вклаадка настроек
theme_frame = tk.LabelFrame(frame_settings, text="Тема оформления", padx=10, pady=10)
theme_frame.pack(fill=tk.X, padx=10, pady=10)

btn_theme = tk.Button(theme_frame, text="Темная тема", command=toggle_theme, width=20)
btn_theme.pack(pady=5)

size_frame = tk.LabelFrame(frame_settings, text="Размер окна", padx=10, pady=10)
size_frame.pack(fill=tk.X, padx=10, pady=10)

btn_fullscreen = tk.Button(size_frame, text="Во весь экран", command=lambda: set_window_size("fullscreen"), width=20)
btn_fullscreen.pack(pady=5)

btn_default = tk.Button(size_frame, text="По умолчанию (1200x900)", command=lambda: set_window_size("default"), width=20)
btn_default.pack(pady=5)
btn_default.config(relief=tk.SUNKEN) 

btn_large = tk.Button(size_frame, text="Большой (1400x1000)", command=lambda: set_window_size("large"), width=20)
btn_large.pack(pady=5)

btn_medium = tk.Button(size_frame, text="Средний (1000x700)", command=lambda: set_window_size("medium"), width=20)
btn_medium.pack(pady=5)

btn_little = tk.Button(size_frame, text="Маленький (700x600)", command=lambda: set_window_size("little"), width=20)
btn_little.pack(pady=5)

apply_theme()

# === Запуск приложения ===
init_db()
update_table_clients()
update_accounts_tree()
root.mainloop()